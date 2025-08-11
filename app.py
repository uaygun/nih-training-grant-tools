import io, re
import numpy as np
import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

# ---------------- UI ----------------
st.set_page_config(page_title="Mentor Tables → 5A/5B DOCX", layout="wide")
st.title("Generate Table 5.docx from Mentor Master Files (Tables 5A & 5B)" ) 

st.markdown("""
**What to upload**
- **Mentor trainee tables (Excel)**: single or multiple `.xlsx` files from mentors. Each file may include sheets named **“Table 5A”** and/or **“Table 5B”**.
- **Mentor publication exports (CSV)**: single or multiple CSV files with publication data for those mentors/trainees.

**What you get**
- One **Word (.docx)** with two sections: **Table 5A** and **Table 5B**.
- Across all uploads, we **merge** sheets, **deduplicate trainees**, **merge mentors** per trainee, and **sort** by **Mentor surname → Trainee surname → Trainee first name**.
""")

st.subheader("1) Upload mentor trainee tables (Excel)")
excel_files = st.file_uploader(
    "Upload mentor Excel files (contain sheets named 'Table 5A' and/or 'Table 5B').",
    type=["xlsx"], accept_multiple_files=True, help="Example: Mentor_A.xlsx, Mentor_B.xlsx"
)

st.subheader("2) Upload mentor publication exports (CSV)")
csv_files = st.file_uploader(
    "Upload publication CSV files corresponding to the mentors.",
    type=["csv"], accept_multiple_files=True, help="Example: Mentor_A_pubs.csv, Mentor_B_pubs.csv"
)

# ------------- Helpers -------------
def bold_name_segments(text, full_name):
    """Return [(segment, is_bold_bool), ...] for bolding trainee in citation."""
    try:
        lastname = full_name.split(',')[0].strip()
        first_part = full_name.split(',')[1].strip()
        first_initial = first_part[0]
    except Exception:
        return [(text, False)]
    pattern = re.compile(rf'\b({re.escape(lastname)}\s*,?\s*{first_initial}[A-Za-z\.]*)\b', re.IGNORECASE)
    matches = list(pattern.finditer(text))
    if not matches:
        return [(text, False)]
    segs, last_idx = [], 0
    for m in matches:
        s, e = m.span()
        if s > last_idx:
            segs.append((text[last_idx:s], False))
        segs.append((m.group(), True))
        last_idx = e
    if last_idx < len(text):
        segs.append((text[last_idx:], False))
    return segs

def build_citation(pub):
    authors = str(pub.get("Authors", "")).strip()
    title   = str(pub.get("Title", "")).strip()
    journal = str(pub.get("Journal/Book", "")).strip()
    date    = str(pub.get("Create Date", "")).strip()
    year    = str(pub.get("Publication Y", "")).strip()
    doi     = str(pub.get("DOI", "")).strip()
    pmid    = str(pub.get("PMID", "")).strip()
    pmcid   = str(pub.get("PMCID", "")).strip()
    try:
        parsed = pd.to_datetime(date)
        date_fmt = parsed.strftime("%Y %b %d")
    except Exception:
        date_fmt = date
    return f"{authors}. {title}. {journal}. {date_fmt};{year}. doi: {doi}. PMID: {pmid}; PMCID: {pmcid}."

def read_sheet_if_exists(xfile, sheet):
    """Read a sheet if present and valid; else return None."""
    try:
        df = pd.read_excel(xfile, sheet_name=sheet)
    except Exception:
        return None
    if df is None or df.empty:
        return None
    df.columns = df.columns.str.strip()
    needed = {"Trainee Name", "Mentor Name"}
    if not needed.issubset(df.columns):
        return None
    return df.dropna(subset=["Trainee Name", "Mentor Name"])

def _normalize_cols(df, cols):
    for c in cols:
        if c in df.columns:
            df[c] = df[c].astype(str).fillna("").str.strip()
        else:
            df[c] = ""
    return df

def merge_sheets(excel_files, sheet_name):
    """
    Merge all instances of a sheet (5A or 5B) across many Excel files.
    Deduplicate by Trainee Name, merge mentors uniquely (joined with ' ; '),
    and sort by Mentor surname -> Trainee surname -> Trainee first name.
    """
    frames = []
    for xf in excel_files:
        df = read_sheet_if_exists(xf, sheet_name)
        if df is not None:
            frames.append(df)
    if not frames:
        return pd.DataFrame(columns=["Trainee Name","Mentor Name","Past/Current","Training Period"])

    merged = pd.concat(frames, ignore_index=True)
    merged = _normalize_cols(merged, ["Trainee Name","Mentor Name","Past/Current","Training Period"])

    # Merge mentors per trainee (unique & sorted). Use ' ; ' to separate mentors.
    mentor_map = (
        merged.groupby("Trainee Name")["Mentor Name"]
        .apply(lambda s: " ; ".join(sorted(set([m.strip() for m in s if m.strip()]))))
        .to_dict()
    )
    # Choose a non-empty training period/status if multiple present
    period_map = (
        merged.groupby("Trainee Name")["Training Period"]
        .apply(lambda s: next((x for x in s if x.strip()), ""))
        .to_dict()
    )
    status_map = (
        merged.groupby("Trainee Name")["Past/Current"]
        .apply(lambda s: next((x for x in s if x.strip()), ""))
        .to_dict()
    )

    uniq = (
        merged.groupby("Trainee Name", as_index=False)
        .agg({"Trainee Name":"first"})
        .rename(columns={"Trainee Name":"Trainee Name"})
    )
    uniq["Mentor Name"]     = uniq["Trainee Name"].map(mentor_map).fillna("")
    uniq["Training Period"] = uniq["Trainee Name"].map(period_map).fillna("")
    uniq["Past/Current"]    = uniq["Trainee Name"].map(status_map).fillna("")

    # Sort keys: Mentor surname (first mentor in list) → Trainee surname → Trainee first
    first_mentor = uniq["Mentor Name"].str.split(" ; ").str[0].fillna("")
    uniq["_mentor_last"] = first_mentor.str.split(",", n=1, expand=True)[0].fillna("").str.strip().str.lower()
    uniq["_trainee_last"]  = uniq["Trainee Name"].str.split(",", n=1, expand=True)[0].fillna("").str.strip().str.lower()
    uniq["_trainee_first"] = uniq["Trainee Name"].str.split(",", n=1, expand=True)[1].fillna("").str.strip().str.lower()

    uniq = uniq.sort_values(by=["_mentor_last","_trainee_last","_trainee_first"]).drop(columns=["_mentor_last","_trainee_last","_trainee_first"])
    return uniq

# ---------- Analytics (UI only; not written to DOCX) ----------
def is_first_author(authors_str: str, trainee_full: str) -> bool:
    """Return True if trainee appears as first author in the Authors string."""
    if not isinstance(authors_str, str) or not isinstance(trainee_full, str):
        return False
    try:
        last = trainee_full.split(",")[0].strip()
        first = trainee_full.split(",")[1].strip()
        initial = first[:1]
    except Exception:
        return False
    pat = re.compile(rf"^\s*{re.escape(last)}\s*,?\s*{re.escape(initial)}(\.|[A-Za-z]\.)?", re.IGNORECASE)
    return bool(pat.search(authors_str))

def section_metrics(section_df: pd.DataFrame, pubs_df: pd.DataFrame):
    """Compute UI-only metrics for a section."""
    if section_df.empty:
        return {"n_trainees":0, "n_with_first":0, "pct_with_first":0.0,
                "total_pubs":0, "min_pubs":0, "median_pubs":0, "max_pubs":0}
    per_counts, with_first, total = [], 0, 0
    for _, row in section_df.iterrows():
        t_name = str(row["Trainee Name"]).strip()
        last = t_name.split(",")[0].strip()
        matches = pubs_df[pubs_df["Authors"].str.contains(last, case=False, na=False)]
        cnt = len(matches)
        total += cnt
        per_counts.append(cnt)
        if any(is_first_author(a, t_name) for a in matches["Authors"].astype(str)):
            with_first += 1
    n = len(section_df)
    pct = (with_first / n * 100.0) if n else 0.0
    arr = np.array(per_counts) if per_counts else np.array([0])
    return {"n_trainees": n, "n_with_first": with_first, "pct_with_first": round(pct,1),
            "total_pubs": int(total), "min_pubs": int(arr.min()), "median_pubs": int(np.median(arr)), "max_pubs": int(arr.max())}

# ------------- DOCX render -------------
def add_section_table(doc, section_df, pubs_df, heading):
    """Add a section where EACH PUBLICATION gets its own table row."""
    if section_df.empty:
        return False
    doc.add_heading(heading, level=1)

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    col_w = [Inches(1.8), Inches(2.2), Inches(1.6), Inches(1.8), Inches(4.0)]
    headers = ["Trainee Name", "Mentor Name", "Past or Current Trainee", "Training Period", "Publication"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]; cell.text = h; cell.width = col_w[i]

    for _, row in section_df.iterrows():
        t_name = str(row["Trainee Name"]).strip()
        mentor = str(row["Mentor Name"]).strip()
        status = str(row.get("Past/Current","")).strip()
        period = str(row.get("Training Period","")).strip()

        last_name = t_name.split(",")[0].strip()
        matches = pubs_df[pubs_df["Authors"].str.contains(last_name, case=False, na=False)]

        # If no pubs, still add a row with empty Publication
        if matches.empty:
            cells = table.add_row().cells
            r = cells[0].paragraphs[0].add_run(t_name); r.bold = True; r.font.size = Pt(10)
            cells[1].text, cells[2].text, cells[3].text = mentor, status, period
            cells[4].text = ""
            continue

        # One row per publication (repeat trainee/mentor info)
        for _, pub in matches.iterrows():
            citation = build_citation(pub)
            cells = table.add_row().cells
            r = cells[0].paragraphs[0].add_run(t_name); r.bold = True; r.font.size = Pt(10)
            cells[1].text, cells[2].text, cells[3].text = mentor, status, period
            cells[4].width = col_w[4]
            p = cells[4].paragraphs[0]
            for seg, is_bold in bold_name_segments(citation, t_name):
                rr = p.add_run(seg); rr.bold = is_bold; rr.font.size = Pt(10)
    return True

# ---------------- Main ----------------
if st.button("Generate DOCX"):
    if not excel_files or not csv_files:
        st.error("Please upload at least one mentor Excel and at least one publication CSV."); st.stop()

    # Merge all publication CSVs
    pubs = pd.DataFrame()
    for f in csv_files:
        pubs = pd.concat([pubs, pd.read_csv(f)], ignore_index=True)

    required = ['Authors', 'Title', 'Journal/Book', 'Create Date', 'DOI', 'PMID', 'PMCID']
    missing = [c for c in required if c not in pubs.columns]
    if missing:
        st.error(f"Missing columns in publication CSVs: {missing}"); st.stop()

    if 'Publication Y' not in pubs.columns:
        pubs['Publication Y'] = pd.to_datetime(pubs['Create Date'], errors='coerce').dt.year.fillna("").astype(str)
    pubs = pubs.dropna(subset=["Authors"])

    # Merge & sort sections
    df_5a = merge_sheets(excel_files, "Table 5A")
    df_5b = merge_sheets(excel_files, "Table 5B")

    # ---------- UI analytics ONLY ----------
    m5a = section_metrics(df_5a, pubs)
    m5b = section_metrics(df_5b, pubs)

    st.subheader("Section Metrics (Shown here only; not included in the Word file)")
    colA, colB = st.columns(2)
    with colA:
        st.markdown("**Table 5A**")
        st.write(f"Trainees: {m5a['n_trainees']}")
        st.write(f"% with ≥1 first-author: {m5a['pct_with_first']}% "
                 f"({m5a['n_with_first']} trainees)")
        st.write(f"Total publications: {m5a['total_pubs']}")
        st.write(f"Pubs per trainee (min/median/max): "
                 f"{m5a['min_pubs']}/{m5a['median_pubs']}/{m5a['max_pubs']}")
    with colB:
        st.markdown("**Table 5B**")
        st.write(f"Trainees: {m5b['n_trainees']}")
        st.write(f"% with ≥1 first-author: {m5b['pct_with_first']}% "
                 f"({m5b['n_with_first']} trainees)")
        st.write(f"Total publications: {m5b['total_pubs']}")
        st.write(f"Pubs per trainee (min/median/max): "
                 f"{m5b['min_pubs']}/{m5b['median_pubs']}/{m5b['max_pubs']}")

    # ---------- Build DOCX (no analytics text) ----------
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5); section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5); section.right_margin  = Inches(0.5)

    any_section = False
    if add_section_table(doc, df_5a, pubs, "Table 5A Publications (Merged by Mentors)"): any_section = True
    if add_section_table(doc, df_5b, pubs, "Table 5B Publications (Merged by Mentors)"): any_section = True
    if not any_section:
        doc.add_heading("No Table 5A or 5B data found in the uploaded mentor files.", level=1)

    out = io.BytesIO()
    doc.save(out); out.seek(0)
    st.success("DOCX ready.")
    st.download_button(
        label="Download DOCX",
        data=out.getvalue(),
        file_name="Table5A_5B_Publications_Merged.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

st.caption("Note: Analytics are displayed here only and are not inserted into the Word file.")
